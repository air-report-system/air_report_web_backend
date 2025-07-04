"""
报告生成服务 - 移植自GUI项目的modify_data_report.py
"""
import os
import re
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Tuple, Optional
from django.conf import settings
from django.core.files.base import ContentFile
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import logging

logger = logging.getLogger(__name__)


class ReportGenerationService:
    """报告生成服务"""
    
    def __init__(self):
        self.template_dir = Path(settings.BASE_DIR) / 'templates' / 'reports'
        self.template_dir.mkdir(parents=True, exist_ok=True)
    
    def generate_report(self, ocr_result_data: Dict[str, Any], form_data: Dict[str, Any]) -> Tuple[bytes, bytes]:
        """
        生成检测报告
        
        Args:
            ocr_result_data: OCR识别结果数据
            form_data: 表单数据
            
        Returns:
            Tuple[bytes, bytes]: (docx文件内容, pdf文件内容)
        """
        try:
            # 准备数据
            report_data = self._prepare_report_data(ocr_result_data, form_data)
            
            # 生成Word文档
            docx_content = self._generate_word_document(report_data)
            
            # 转换为PDF
            pdf_content = self._convert_to_pdf(docx_content)
            
            return docx_content, pdf_content
            
        except Exception as e:
            logger.error(f"报告生成失败: {e}")
            raise e
    
    def _prepare_report_data(self, ocr_result_data: Dict[str, Any], form_data: Dict[str, Any]) -> Dict[str, Any]:
        """准备报告数据"""
        # 从OCR结果提取点位数据
        points_data = ocr_result_data.get('points_data', {})
        
        # 转换点位数据格式
        point_list = []
        for point_name, point_value in points_data.items():
            try:
                value_float = float(point_value)
                point_list.append((point_name, f"{value_float:.3f}"))
            except (ValueError, TypeError):
                point_list.append((point_name, str(point_value)))
        
        # 判断检测类型
        check_type = self._determine_check_type(points_data, form_data.get('check_type'))
        
        # 准备日期信息
        date_info = self._prepare_date_info(form_data.get('sampling_date', ''))
        
        return {
            'project_address': form_data.get('project_address', ''),
            'contact_person': form_data.get('contact_person', ''),
            'phone': ocr_result_data.get('phone', ''),
            'sampling_date': form_data.get('sampling_date', ''),
            'temperature': form_data.get('temperature', ocr_result_data.get('temperature', '')),
            'humidity': form_data.get('humidity', ocr_result_data.get('humidity', '')),
            'check_type': check_type,
            'check_type_display': '初检' if check_type == 'initial' else '复检',
            'points_data': point_list,
            'date_info': date_info
        }
    
    def _determine_check_type(self, points_data: Dict[str, Any], form_check_type: Optional[str] = None) -> str:
        """
        判断检测类型 - 移植自GUI项目的逻辑
        基于点位值的众数进行判断
        """
        if form_check_type:
            return form_check_type
        
        if not points_data:
            return 'initial'
        
        try:
            # 统计>0.080和≤0.080的点位数量
            high_count = 0  # >0.080
            low_count = 0   # ≤0.080
            
            for point_value in points_data.values():
                try:
                    value = float(point_value)
                    if value > 0.080:
                        high_count += 1
                    else:
                        low_count += 1
                except (ValueError, TypeError):
                    continue
            
            # 根据众数判断
            if high_count > low_count:
                return 'recheck'  # 复检
            else:
                return 'initial'  # 初检
                
        except Exception as e:
            logger.warning(f"检测类型判断失败，使用默认值: {e}")
            return 'initial'
    
    def _prepare_date_info(self, sampling_date: str) -> Dict[str, str]:
        """准备日期信息"""
        logger.info(f"[DEBUG] _prepare_date_info 输入: sampling_date='{sampling_date}'")
        try:
            if sampling_date:
                # 尝试解析日期
                if '-' in sampling_date:
                    parts = sampling_date.split('-')
                    logger.info(f"[DEBUG] 分割日期: parts={parts}")

                    if len(parts) == 3:
                        # 完整日期格式: YYYY-MM-DD
                        year, month, day = parts
                        result = {'month': month.zfill(2), 'day': day.zfill(2)}
                        logger.info(f"[DEBUG] 完整日期格式，提取月日: {result}")
                        return result
                    elif len(parts) == 2:
                        # 月日格式: MM-DD
                        month, day = parts
                        result = {'month': month.zfill(2), 'day': day.zfill(2)}
                        logger.info(f"[DEBUG] 月日格式: {result}")
                        return result

            # 使用当前日期
            now = datetime.now()
            result = {
                'month': f"{now.month:02d}",
                'day': f"{now.day:02d}"
            }
            logger.info(f"[DEBUG] 使用当前日期: {result}")
            return result

        except Exception as e:
            logger.warning(f"[DEBUG] _prepare_date_info 异常: {e}")
            now = datetime.now()
            result = {
                'month': f"{now.month:02d}",
                'day': f"{now.day:02d}"
            }
            logger.info(f"[DEBUG] 异常后使用当前日期: {result}")
            return result
    
    def _generate_word_document(self, report_data: Dict[str, Any]) -> bytes:
        """
        生成Word文档 - 移植自GUI项目的modify_data_report.py
        """
        try:
            # 获取模板文件
            template_path = self._get_template_path()
            
            if template_path and template_path.exists():
                doc = Document(str(template_path))
            else:
                # 创建新文档
                doc = Document()
                self._create_default_template(doc)
            
            # 替换文档中的占位符
            self._replace_placeholders(doc, report_data)
            
            # 处理表格
            self._process_tables(doc, report_data)
            
            # 保存到内存
            from io import BytesIO
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            return doc_buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Word文档生成失败: {e}")
            raise e
    
    def _get_template_path(self) -> Optional[Path]:
        """获取模板文件路径"""
        template_files = [
            '模板.docx',  # 添加用户提供的模板文件名
            'report_template.docx',
            'template.docx',
            '检测报告模板.docx'
        ]

        for filename in template_files:
            template_path = self.template_dir / filename
            if template_path.exists():
                logger.info(f"找到模板文件: {template_path}")
                return template_path

        logger.warning(f"未找到模板文件，搜索路径: {self.template_dir}")
        return None



    def _get_compatible_font(self, original_font: str) -> str:
        """获取与原始字体兼容的系统字体"""
        font_mapping = {
            # 中文字体
            '宋体': 'SimSun',
            'SimSun': 'SimSun',
            '新宋体': 'SimSun',
            'NSimSun': 'SimSun',
            '仿宋': 'FangSong',
            'FangSong': 'FangSong',
            '仿宋_GB2312': 'FangSong',
            '黑体': 'SimHei',
            'SimHei': 'SimHei',
            '微软雅黑': 'SimHei',  # 用黑体替代微软雅黑
            'Microsoft YaHei': 'SimHei',

            # 英文字体
            'Arial': 'Arial',
            'Arial Black': 'Arial',
            'Times New Roman': 'Times-Roman',
            'Times': 'Times-Roman',
            'Calibri': 'Calibri',

            # 其他常见字体映射
            'Helvetica': 'Arial',
            'Verdana': 'Arial',
            'Tahoma': 'Arial',
            'Georgia': 'Times-Roman',
        }

        return font_mapping.get(original_font, 'SimSun')  # 默认使用宋体

    def _set_chinese_font(self, paragraph, preserve_original: bool = True):
        """设置段落的字体，严格保持原始字体设置"""
        try:
            from docx.shared import Pt
            from docx.oxml.ns import qn

            for run in paragraph.runs:
                if run.text.strip():
                    # 获取原始字体设置
                    original_ascii_font = None
                    original_eastasia_font = None

                    if preserve_original:
                        r = run._element
                        rPr = r.find(qn('w:rPr'))
                        if rPr is not None:
                            rFonts = rPr.find(qn('w:rFonts'))
                            if rFonts is not None:
                                original_ascii_font = rFonts.get(qn('w:ascii'))
                                original_eastasia_font = rFonts.get(qn('w:eastAsia'))

                    # 设置东亚字体（对中文很重要）
                    r = run._element
                    rPr = r.get_or_add_rPr()
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = rPr.makeelement(qn('w:rFonts'))
                        rPr.append(rFonts)

                    # 只更新中文字体，保持英文字体不变
                    if original_eastasia_font:
                        # 如果有原始东亚字体，使用兼容字体替换
                        target_eastasia_font = self._get_compatible_font(original_eastasia_font)
                        rFonts.set(qn('w:eastAsia'), target_eastasia_font)
                    else:
                        # 如果没有原始东亚字体，设置默认中文字体
                        rFonts.set(qn('w:eastAsia'), 'SimSun')

                    # 保持原始的ASCII字体（英文字体）不变
                    if original_ascii_font:
                        # 如果原始字体是中文字体，才进行替换
                        if original_ascii_font in ['宋体', 'SimSun', '黑体', 'SimHei', '仿宋', 'FangSong']:
                            target_ascii_font = self._get_compatible_font(original_ascii_font)
                            rFonts.set(qn('w:ascii'), target_ascii_font)
                            rFonts.set(qn('w:hAnsi'), target_ascii_font)
                        # 如果是英文字体，保持不变
                        else:
                            rFonts.set(qn('w:ascii'), original_ascii_font)
                            rFonts.set(qn('w:hAnsi'), original_ascii_font)
                    # 如果原始ASCII字体为None，不设置，保持默认

        except Exception as e:
            logger.warning(f"设置字体失败: {e}")

    def _fix_document_formatting(self, doc: Document):
        """修复文档格式问题"""
        try:
            # 不修改声明的对齐，保持原始模板格式
            # 只进行必要的字体修复
            pass

        except Exception as e:
            logger.warning(f"修复文档格式失败: {e}")

    def _create_default_template(self, doc: Document):
        """创建默认模板"""
        # 添加标题
        title = doc.add_heading('室内空气质量检测报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加基本信息段落
        doc.add_paragraph('项目地址：{project_address}')
        doc.add_paragraph('联系人：{contact_person}')
        doc.add_paragraph('联系电话：{phone}')
        doc.add_paragraph('采样日期：{sampling_date}')
        doc.add_paragraph('现场温度：{temperature}℃')
        doc.add_paragraph('现场湿度：{humidity}%')
        doc.add_paragraph('检测类型：{check_type_display}')
        
        # 添加检测结果表格占位符
        doc.add_paragraph('检测结果：')
        doc.add_paragraph('[检测结果表格]')
    
    def _replace_placeholders(self, doc: Document, report_data: Dict[str, Any]):
        """替换文档中的占位符 - 完全移植自GUI项目的替换逻辑"""
        # 获取日期信息
        logger.info(f"[DEBUG] _replace_placeholders 开始处理日期")
        date_info = report_data.get('date_info', {})
        logger.info(f"[DEBUG] 从report_data获取的date_info: {date_info}")
        month = date_info.get('month', '')
        day = date_info.get('day', '')
        logger.info(f"[DEBUG] 初始 month='{month}', day='{day}'")

        # 如果date_info中没有数据，尝试从sampling_date解析
        if not month or not day:
            sampling_date = report_data.get('sampling_date', '')
            logger.info(f"[DEBUG] date_info为空，尝试解析sampling_date: '{sampling_date}'")
            if sampling_date:
                try:
                    from datetime import datetime
                    if '-' in sampling_date:
                        date_obj = datetime.strptime(sampling_date, "%m-%d")
                        month = str(date_obj.month)
                        day = str(date_obj.day)
                        logger.info(f"[DEBUG] 从sampling_date解析得到 month='{month}', day='{day}'")
                    elif '/' in sampling_date:
                        date_obj = datetime.strptime(sampling_date, "%m/%d")
                        month = str(date_obj.month)
                        day = str(date_obj.day)
                        logger.info(f"[DEBUG] 从sampling_date解析得到 month='{month}', day='{day}'")
                except Exception as e:
                    logger.warning(f"[DEBUG] 解析sampling_date失败: {e}")

        # 确保月日是字符串格式（去掉前导零）
        if month:
            original_month = month
            month = str(int(month))  # 去掉前导零，如 "07" -> "7"
            logger.info(f"[DEBUG] 月份去前导零: '{original_month}' -> '{month}'")
        if day:
            original_day = day
            day = str(int(day))      # 去掉前导零，如 "03" -> "3"
            logger.info(f"[DEBUG] 日期去前导零: '{original_day}' -> '{day}'")

        logger.info(f"[DEBUG] 最终用于替换的 month='{month}', day='{day}'")

        # 修改第一页的内容 - 完全按照GUI版本的逻辑
        logger.info(f"[DEBUG] 开始处理段落替换，共{len(doc.paragraphs)}个段落")
        for para_idx, paragraph in enumerate(doc.paragraphs):
            content_changed = False
            if '<月>' in paragraph.text or '<日>' in paragraph.text:
                logger.info(f"[DEBUG] 段落{para_idx}包含日期占位符: '{paragraph.text}'")

            for run_idx, run in enumerate(paragraph.runs):
                original_text = run.text
                new_text = original_text

                # 替换占位符
                if "<地址>" in new_text:
                    replacement = report_data.get('project_address', '')
                    new_text = new_text.replace("<地址>", replacement)
                    logger.info(f"[DEBUG] 段落{para_idx}-run{run_idx}: 替换<地址> -> '{replacement}'")
                    content_changed = True
                if "<联系人>" in new_text:
                    replacement = report_data.get('contact_person', '')
                    new_text = new_text.replace("<联系人>", replacement)
                    logger.info(f"[DEBUG] 段落{para_idx}-run{run_idx}: 替换<联系人> -> '{replacement}'")
                    content_changed = True
                if "<月>" in new_text and month:
                    new_text = new_text.replace("<月>", month)
                    logger.info(f"[DEBUG] 段落{para_idx}-run{run_idx}: 替换<月> -> '{month}', 原文: '{original_text}' -> 新文: '{new_text}'")
                    content_changed = True
                if "<日>" in new_text and day:
                    new_text = new_text.replace("<日>", day)
                    logger.info(f"[DEBUG] 段落{para_idx}-run{run_idx}: 替换<日> -> '{day}', 原文: '{original_text}' -> 新文: '{new_text}'")
                    content_changed = True
                if "<温度>" in new_text:
                    replacement = report_data.get('temperature', '')
                    new_text = new_text.replace("<温度>", replacement)
                    logger.info(f"[DEBUG] 段落{para_idx}-run{run_idx}: 替换<温度> -> '{replacement}'")
                    content_changed = True
                if "<湿度>" in new_text:
                    replacement = report_data.get('humidity', '')
                    new_text = new_text.replace("<湿度>", replacement)
                    logger.info(f"[DEBUG] 段落{para_idx}-run{run_idx}: 替换<湿度> -> '{replacement}'")
                    content_changed = True

                # 如果有变化，更新文本但保持格式
                if new_text != original_text:
                    run.text = new_text
                    content_changed = True

            # 如果内容有变化，设置兼容的中文字体（保持原始字体风格）
            if content_changed:
                logger.info(f"[DEBUG] 段落{para_idx}内容已更改，设置字体")
                self._set_chinese_font(paragraph, preserve_original=True)

        # 修复特定格式问题
        self._fix_document_formatting(doc)

        # 修改第三页委托概况的内容 - 完全按照GUI版本的逻辑
        logger.info(f"[DEBUG] 开始处理表格替换，共{len(doc.tables)}个表格")
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    content_changed = False
                    original_text = cell.text
                    new_text = original_text

                    if '<月>' in original_text or '<日>' in original_text:
                        logger.info(f"[DEBUG] 表格{table_idx}-行{row_idx}-列{cell_idx}包含日期占位符: '{original_text}'")

                    # 替换各种占位符
                    if "<联系人>" in new_text:
                        replacement = report_data.get('contact_person', '')
                        new_text = new_text.replace("<联系人>", replacement)
                        logger.info(f"[DEBUG] 表格{table_idx}-行{row_idx}-列{cell_idx}: 替换<联系人> -> '{replacement}'")
                        content_changed = True
                    if "<地址>" in new_text:
                        replacement = report_data.get('project_address', '')
                        new_text = new_text.replace("<地址>", replacement)
                        logger.info(f"[DEBUG] 表格{table_idx}-行{row_idx}-列{cell_idx}: 替换<地址> -> '{replacement}'")
                        content_changed = True
                    if "<月>" in new_text and month:
                        new_text = new_text.replace("<月>", month)
                        logger.info(f"[DEBUG] 表格{table_idx}-行{row_idx}-列{cell_idx}: 替换<月> -> '{month}', 原文: '{original_text}' -> 新文: '{new_text}'")
                        content_changed = True
                    if "<日>" in new_text and day:
                        new_text = new_text.replace("<日>", day)
                        logger.info(f"[DEBUG] 表格{table_idx}-行{row_idx}-列{cell_idx}: 替换<日> -> '{day}', 原文: '{original_text}' -> 新文: '{new_text}'")
                        content_changed = True
                    if "<温度>" in new_text:
                        replacement = report_data.get('temperature', '')
                        new_text = new_text.replace("<温度>", replacement)
                        logger.info(f"[DEBUG] 表格{table_idx}-行{row_idx}-列{cell_idx}: 替换<温度> -> '{replacement}'")
                        content_changed = True
                    if "<湿度>" in new_text:
                        replacement = report_data.get('humidity', '')
                        new_text = new_text.replace("<湿度>", replacement)
                        logger.info(f"[DEBUG] 表格{table_idx}-行{row_idx}-列{cell_idx}: 替换<湿度> -> '{replacement}'")
                        content_changed = True

                    # 如果有变化，更新单元格文本
                    if new_text != original_text:
                        cell.text = new_text
                        logger.info(f"[DEBUG] 表格{table_idx}-行{row_idx}-列{cell_idx}: 文本已更新")
                        content_changed = True

                    # 如果内容有变化，设置兼容的中文字体（保持原始字体风格）
                    if content_changed:
                        for paragraph in cell.paragraphs:
                            self._set_chinese_font(paragraph, preserve_original=True)

        # 修改环境条件的内容 - 完全按照GUI版本的逻辑
        for paragraph in doc.paragraphs:
            content_changed = False
            if "<温度>" in paragraph.text:
                paragraph.text = paragraph.text.replace("<温度>", report_data.get('temperature', ''))
                content_changed = True
            if "<湿度>" in paragraph.text:
                paragraph.text = paragraph.text.replace("<湿度>", report_data.get('humidity', ''))
                content_changed = True

            # 如果内容有变化，设置兼容的中文字体（保持原始字体风格）
            if content_changed:
                self._set_chinese_font(paragraph, preserve_original=True)

    
    def _process_tables(self, doc: Document, report_data: Dict[str, Any]):
        """
        处理表格 - 完全移植自GUI项目的表格处理逻辑
        """
        points_data = report_data.get('points_data', [])
        
        if not points_data:
            logger.warning("没有点位数据，跳过表格处理")
            return
        
        # 处理检测结果表格 - 完全按照GUI版本的逻辑
        for table_idx, table in enumerate(doc.tables):
            # 尝试检测是否是目标表格 - 按照GUI版本的逻辑
            try:
                header_text = table.cell(0, 0).text.strip()
                is_target_table = False
                if ("序号" in header_text or "序" == header_text) and (
                    "Number" in header_text
                ):
                    is_target_table = True
                elif len(table.rows) > 1 and len(table.columns) > 1:
                    second_cell_text = table.cell(0, 1).text.strip()
                    if "检测位置" in second_cell_text or "Analysis" in second_cell_text:
                        is_target_table = True
            except:
                continue

            if is_target_table:
                logger.info(f"找到检测结果表格（索引：{table_idx}），共有{len(points_data)}个点位数据")
                
                # 创建新表格替换原表格 - 完全按照GUI版本的逻辑
                self._create_results_table_gui_style(doc, table, points_data)
                break
    
    def _create_results_table_gui_style(self, doc: Document, original_table, points_data: List[Tuple[str, str]]):
        """
        创建检测结果表格 - 完全按照GUI项目逻辑实现
        """
        try:
            # 获取表格位置
            parent = original_table._element.getparent()
            index = parent.index(original_table._element)
            
            # 创建新表格（4列）
            # 行数 = 3(表头) + 数据行数（至少4行） + 2(备注行)
            rows_needed = 3 + max(4, len(points_data)) + 2
            logger.info(f"创建新表格，总行数：{rows_needed}（3行表头 + {max(4, len(points_data))}行数据 + 2行备注）")
            new_table = doc.add_table(rows=rows_needed, cols=4)

            # 设置表格样式
            if hasattr(original_table, 'style'):
                new_table.style = original_table.style
                logger.info(f"应用原表格样式")

            # 导入XML处理模块
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            # 设置单元格边框函数 - 完全按照GUI版本
            def set_cell_border(cell, border_type="all", size=4):
                """设置单元格边框"""
                tc = cell._element.tcPr
                if tc is None:
                    tc = OxmlElement("w:tcPr")
                    cell._element.append(tc)

                # 清除现有边框
                for border in tc.findall(qn("w:tcBorders")):
                    tc.remove(border)

                borders = OxmlElement("w:tcBorders")

                if border_type == "all" or border_type == "top":
                    top = OxmlElement("w:top")
                    top.set(qn("w:val"), "single")
                    top.set(qn("w:sz"), str(size))
                    top.set(qn("w:space"), "0")
                    top.set(qn("w:color"), "auto")
                    borders.append(top)

                if border_type == "all" or border_type == "bottom":
                    bottom = OxmlElement("w:bottom")
                    bottom.set(qn("w:val"), "single")
                    bottom.set(qn("w:sz"), str(size))
                    bottom.set(qn("w:space"), "0")
                    bottom.set(qn("w:color"), "auto")
                    borders.append(bottom)

                if border_type == "all" or border_type == "left":
                    left = OxmlElement("w:left")
                    left.set(qn("w:val"), "single")
                    left.set(qn("w:sz"), str(size))
                    left.set(qn("w:space"), "0")
                    left.set(qn("w:color"), "auto")
                    borders.append(left)

                if border_type == "all" or border_type == "right":
                    right = OxmlElement("w:right")
                    right.set(qn("w:val"), "single")
                    right.set(qn("w:sz"), str(size))
                    right.set(qn("w:space"), "0")
                    right.set(qn("w:color"), "auto")
                    borders.append(right)

                tc.append(borders)

                # 设置单元格垂直居中
                tcVAlign = OxmlElement("w:vAlign")
                tcVAlign.set(qn("w:val"), "center")
                tc.append(tcVAlign)

            # 设置行高的函数 - 完全按照GUI版本
            def set_row_height(row, height):
                """设置行高"""
                tr = row._tr
                trPr = tr.get_or_add_trPr()
                trHeight = OxmlElement("w:trHeight")
                trHeight.set(qn("w:val"), str(height))
                trHeight.set(qn("w:hRule"), "atLeast")

                # 删除旧的行高设置
                for old_height in trPr.findall(qn("w:trHeight")):
                    trPr.remove(old_height)

                trPr.append(trHeight)

            # 设置单元格宽度的函数 - 完全按照GUI版本
            def set_column_width(table, col_index, width):
                """设置列宽"""
                for cell in table.columns[col_index].cells:
                    tc = cell._element.tcPr
                    if tc is None:
                        tc = OxmlElement("w:tcPr")
                        cell._element.append(tc)

                    tcW = OxmlElement("w:tcW")
                    tcW.set(qn("w:w"), str(width))
                    tcW.set(qn("w:type"), "dxa")

                    # 移除旧的宽度设置
                    for old_width in tc.findall(qn("w:tcW")):
                        tc.remove(old_width)

                    tc.append(tcW)

            # 辅助函数：设置单元格文本和样式 - 完全按照GUI版本
            def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
                cell.text = text
                for paragraph in cell.paragraphs:
                    paragraph.alignment = align
                    # 设置段落间距
                    paragraph.space_before = Pt(2)
                    paragraph.space_after = Pt(2)
                    for run in paragraph.runs:
                        run.font.size = Pt(12)  # 使用小四号字体(约12pt)
                        run.bold = bold
                    # 设置兼容的中文字体（保持原始字体风格）
                    self._set_chinese_font(paragraph, preserve_original=False)  # 新创建的表格使用默认字体
                # 添加边框
                set_cell_border(cell)

            # 为表头设置灰色背景 - 完全按照GUI版本
            def set_cell_shading(cell, color="F2F2F2"):
                """设置单元格背景颜色"""
                tc = cell._element.tcPr
                if tc is None:
                    tc = OxmlElement("w:tcPr")
                    cell._element.append(tc)

                # 删除现有的shading
                for shading in tc.findall(qn("w:shd")):
                    tc.remove(shading)

                # 添加新的shading
                shading = OxmlElement("w:shd")
                shading.set(qn("w:val"), "clear")
                shading.set(qn("w:color"), "auto")
                shading.set(qn("w:fill"), color)
                tc.append(shading)

            # 为所有单元格添加边框，设置行高，并设置列宽 - 完全按照GUI版本
            for row in new_table.rows:
                # 设置行高约为小四号字体高度的1.5倍
                set_row_height(row, 360)  # 约18pt
                for cell in row.cells:
                    set_cell_border(cell)

            # 设置列宽 - 完全按照GUI版本
            from docx.shared import Inches
            total_width = Inches(6.5)  # 假设总宽度约为6.5英寸
            set_column_width(new_table, 0, int(total_width.emu * 0.15))  # 序号列宽度为15%
            set_column_width(new_table, 1, int(total_width.emu * 0.4))   # 检测位置列宽度为40%
            set_column_width(new_table, 2, int(total_width.emu * 0.225)) # 检测值列宽度为22.5%
            set_column_width(new_table, 3, int(total_width.emu * 0.225)) # 限值列宽度为22.5%

            logger.info("表格边框和行高设置完成")

            # 按照GUI版本设置表头和合并单元格
            # 第一行表头
            set_cell_text(new_table.cell(0, 0), "序号", True)
            set_cell_text(new_table.cell(0, 1), "检测位置\nAnalysistems", True)
            set_cell_text(new_table.cell(0, 2), "检测值Test Results", True)
            new_table.cell(0, 3).text = ""  # 先清空内容再合并

            # 第二行表头 - 设置灰色背景
            set_cell_text(new_table.cell(1, 2), "甲醛 (mg/m³)", True)
            new_table.cell(1, 3).text = ""  # 先清空内容再合并

            # 第三行表头
            set_cell_text(new_table.cell(2, 2), "检测值", True)
            set_cell_text(new_table.cell(2, 3), "限值", True)

            # 按照GUI版本进行单元格合并
            new_table.cell(0, 0).merge(new_table.cell(2, 0))  # 序号跨3行
            new_table.cell(0, 1).merge(new_table.cell(2, 1))  # 检测位置跨3行
            new_table.cell(0, 2).merge(new_table.cell(0, 3))  # 检测值Test Results跨2列
            new_table.cell(1, 2).merge(new_table.cell(1, 3))  # 甲醛跨2列

            # 为所有表头单元格设置灰色背景 - 完全按照GUI版本
            for i in range(3):
                for j in range(4):
                    try:
                        cell = new_table.cell(i, j)
                        set_cell_shading(cell)
                    except IndexError:
                        # 跳过已合并的单元格
                        pass

            logger.info("表头设置和单元格合并完成")

            # 填充数据行 - 完全按照GUI版本
            data_start_row = 3  # 数据从第4行开始(索引为3)
            logger.info(f"开始填充数据行，共{len(points_data)}条记录")

            for i, (point, value) in enumerate(points_data):
                row_idx = data_start_row + i
                if row_idx >= len(new_table.rows) - 2:  # 预留2行给备注
                    logger.warning(f"数据行 {i + 1} 超出表格范围，跳过")
                    continue

                try:
                    # 序号
                    set_cell_text(new_table.cell(row_idx, 0), str(i + 1))
                    # 点位
                    set_cell_text(new_table.cell(row_idx, 1), point)
                    # 值
                    set_cell_text(new_table.cell(row_idx, 2), value)
                    # 限值
                    set_cell_text(new_table.cell(row_idx, 3), "≤0.08", False)
                except Exception as e:
                    logger.warning(f"填充数据行{i + 1}异常: {str(e)}")

            # 如果点位数据少于4个，填充剩余空行 - 完全按照GUI版本
            if len(points_data) < 4:
                logger.info(f"填充剩余空行（{4 - len(points_data)}行）")
                for i in range(len(points_data), 4):
                    row_idx = data_start_row + i
                    # 序号
                    set_cell_text(new_table.cell(row_idx, 0), str(i + 1))
                    # 其他单元格保持空白但添加边框
                    for col in range(1, 4):
                        set_cell_border(new_table.cell(row_idx, col))

            # 添加表格内的备注行 - 完全按照GUI版本
            footer_start_row = data_start_row + max(4, len(points_data))
            if footer_start_row + 1 < len(new_table.rows):
                # 备注1
                remark1_row = new_table.row_cells(footer_start_row)
                remark1_row[0].merge(remark1_row[1])  # 合并前两列
                remark1_row[2].merge(remark1_row[3])  # 合并后两列

                set_cell_text(remark1_row[0], "备注Remarks ①", False, WD_ALIGN_PARAGRAPH.LEFT)
                set_cell_text(remark1_row[2], "分析依据：《室内空气质量标准》 GB/T 18883-2022", False, WD_ALIGN_PARAGRAPH.LEFT)

                # 备注2
                remark2_row = new_table.row_cells(footer_start_row + 1)
                remark2_row[0].merge(remark2_row[1])  # 合并前两列
                remark2_row[2].merge(remark2_row[3])  # 合并后两列

                set_cell_text(remark2_row[0], "备注Remarks ②", False, WD_ALIGN_PARAGRAPH.LEFT)
                set_cell_text(remark2_row[2], "GB/T 18883-2002部分条款已被GB/T 18883-2022替代", False, WD_ALIGN_PARAGRAPH.LEFT)

                logger.info("表格备注行添加完成")

            # 替换原表格
            parent[index] = new_table._element
            logger.info("原表格已替换为新表格")
            logger.info("表格创建完成")
            
        except Exception as e:
            logger.error(f"创建检测结果表格失败: {e}")
            raise e
    
    def _convert_to_pdf(self, docx_content: bytes) -> bytes:
        """
        将Word文档转换为PDF - 只使用能保持完整格式的转换方法
        
        绝对不使用任何格式简化的降级方案。
        格式完整性是报告质量的基本要求。
        """
        try:
            # 创建临时文件
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
                temp_docx.write(docx_content)
                temp_docx_path = temp_docx.name

            try:
                # 验证docx文件格式完整性
                self._validate_docx_format(temp_docx_path)
                
                # 生成PDF文件路径
                temp_pdf_path = temp_docx_path.replace('.docx', '.pdf')

                # 只尝试格式保持的PDF转换方法
                pdf_content = self._try_pdf_conversion_methods(temp_docx_path, temp_pdf_path)

                return pdf_content

            finally:
                # 清理临时文件
                for temp_file in [temp_docx_path, temp_pdf_path]:
                    if os.path.exists(temp_file):
                        try:
                            os.unlink(temp_file)
                        except Exception:
                            pass

        except Exception as e:
            logger.error(f"[DEBUG] PDF转换失败: {e}")
            # 绝对不提供降级方案，直接抛出异常
            raise RuntimeError(f"PDF转换失败，无法保证格式完整性: {e}")

    def _try_pdf_conversion_methods(self, docx_path: str, pdf_path: str) -> bytes:
        """尝试不同的PDF转换方法"""
        logger.info(f"[DEBUG] 开始PDF转换，docx路径: {docx_path}")
        logger.info(f"[DEBUG] 目标PDF路径: {pdf_path}")

        # 方法1: 使用unoconv转换（推荐方法）
        logger.info(f"[DEBUG] 尝试方法1: unoconv转换")
        try:
            import subprocess
            import os

            # 设置环境变量以支持中文字符和UNO路径
            env = os.environ.copy()
            env['LC_ALL'] = 'zh_CN.UTF-8'
            env['LANG'] = 'zh_CN.UTF-8'
            env['UNO_PATH'] = '/usr/lib/libreoffice/program'
            env['PYTHONPATH'] = '/usr/lib/python3/dist-packages:/usr/lib/libreoffice/program'

            # 使用unoconv转换，指定使用系统python3
            command_args = [
                '/usr/bin/python3', '/usr/bin/unoconv',
                '-f', 'pdf',
                '-o', pdf_path,
                docx_path
            ]

            logger.info(f"[DEBUG] 执行unoconv命令: {' '.join(command_args)}")

            result = subprocess.run(
                command_args,
                capture_output=True,
                timeout=120,  # 增加超时时间到2分钟
                text=True,
                env=env
            )

            logger.info(f"[DEBUG] unoconv命令执行完成，返回码: {result.returncode}")
            logger.info(f"[DEBUG] 标准输出: {result.stdout}")
            if result.stderr:
                logger.info(f"[DEBUG] 标准错误: {result.stderr}")

            # 检查PDF文件是否存在且有效
            if result.returncode == 0 and os.path.exists(pdf_path):
                pdf_size = os.path.getsize(pdf_path)
                logger.info(f"[DEBUG] PDF文件生成成功，大小: {pdf_size} bytes")

                if pdf_size > 1000:  # 至少1KB
                    with open(pdf_path, 'rb') as pdf_file:
                        pdf_content = pdf_file.read()
                        # 验证是否是有效的PDF文件
                        if pdf_content.startswith(b'%PDF'):
                            logger.info("[DEBUG] unoconv转换成功，返回PDF内容")
                            return pdf_content
                        else:
                            logger.warning("[DEBUG] unoconv生成的文件不是有效的PDF")
                else:
                    logger.warning(f"[DEBUG] unoconv生成的PDF文件太小: {pdf_size} bytes")
            else:
                logger.warning(f"[DEBUG] unoconv转换失败，返回码: {result.returncode}")
                logger.warning(f"[DEBUG] PDF文件存在: {os.path.exists(pdf_path)}")

        except FileNotFoundError:
            logger.warning("[DEBUG] unoconv命令未找到，请确保已安装unoconv")
        except subprocess.TimeoutExpired:
            logger.warning("[DEBUG] unoconv转换超时")
        except Exception as e:
            logger.warning(f"[DEBUG] unoconv转换失败: {e}")

        # 方法2: 使用LibreOffice命令行（备选方案）
        logger.info(f"[DEBUG] 尝试方法2: LibreOffice转换")
        try:
            # 尝试不同的LibreOffice命令名称
            libreoffice_commands = [
                'libreoffice',  # 标准命令
                'soffice',      # 备选命令
                '/usr/bin/libreoffice',  # Ubuntu/Debian标准路径
                '/usr/bin/soffice',      # 备选路径
                '/opt/libreoffice/program/soffice',  # 自定义安装路径
                '/snap/bin/libreoffice', # Snap安装路径
                '/usr/local/bin/libreoffice',  # 手动编译安装路径
            ]

            for cmd in libreoffice_commands:
                logger.info(f"[DEBUG] 尝试LibreOffice命令: {cmd}")
                try:
                    # 设置环境变量支持中文
                    env = os.environ.copy()
                    env['LC_ALL'] = 'zh_CN.UTF-8'
                    env['LANG'] = 'zh_CN.UTF-8'

                    # 使用LibreOffice转换参数
                    command_args = [
                        cmd,
                        '--headless',
                        '--invisible',
                        '--nodefault',
                        '--nolockcheck',
                        '--nologo',
                        '--norestore',
                        '--convert-to', 'pdf:writer_pdf_Export',
                        '--outdir', os.path.dirname(pdf_path),
                        docx_path
                    ]
                    logger.info(f"[DEBUG] 执行LibreOffice命令: {' '.join(command_args)}")

                    result = subprocess.run(command_args, capture_output=True, timeout=120, text=True, env=env)

                    logger.info(f"[DEBUG] LibreOffice命令执行完成，返回码: {result.returncode}")
                    logger.info(f"[DEBUG] 标准输出: {result.stdout}")
                    if result.stderr:
                        logger.info(f"[DEBUG] 标准错误: {result.stderr}")

                    if result.returncode == 0 and os.path.exists(pdf_path):
                        pdf_size = os.path.getsize(pdf_path)
                        logger.info(f"[DEBUG] LibreOffice PDF文件生成成功，大小: {pdf_size} bytes")

                        if pdf_size > 1000:  # 至少1KB
                            with open(pdf_path, 'rb') as pdf_file:
                                pdf_content = pdf_file.read()
                                if pdf_content.startswith(b'%PDF'):
                                    logger.info(f"[DEBUG] LibreOffice PDF转换成功，使用命令: {cmd}")
                                    return pdf_content
                                else:
                                    logger.warning("[DEBUG] LibreOffice生成的文件不是有效的PDF")
                        else:
                            logger.warning(f"[DEBUG] LibreOffice生成的PDF文件太小: {pdf_size} bytes")
                    else:
                        logger.warning(f"[DEBUG] LibreOffice转换失败，命令: {cmd}, 返回码: {result.returncode}")

                except FileNotFoundError:
                    continue
                except subprocess.TimeoutExpired:
                    logger.warning(f"[DEBUG] LibreOffice转换超时: {cmd}")
                    continue

        except Exception as e:
            logger.warning(f"LibreOffice转换异常: {e}")

        # 如果所有格式保持方法都失败，直接抛出异常
        error_msg = """PDF转换失败！

所有能够保持文档格式的转换方法都不可用。

请确保安装以下任一工具：
1. unoconv (推荐): sudo apt-get install unoconv
2. LibreOffice: sudo apt-get install libreoffice

unoconv方案相比其他方法在测试中表现更好，能够更好地保持文档格式。

安装方法：
Ubuntu/Debian: sudo apt-get install unoconv libreoffice
CentOS/RHEL: sudo yum install unoconv libreoffice
或运行项目提供的安装脚本: backend/scripts/install_pdf_dependencies.sh

注意：为保证报告格式的完整性和准确性，系统不提供任何格式简化的降级方案。
格式完整性是报告质量的基本要求。"""

        logger.error(f"[DEBUG] 所有格式保持转换方法都失败，拒绝生成格式错误的PDF")
        logger.error(error_msg)
        raise RuntimeError(error_msg)

    def _validate_docx_format(self, docx_path: str):
        """验证docx文件格式完整性"""
        try:
            from docx import Document
            
            # 验证文件可以被正确读取
            doc = Document(docx_path)
            logger.info(f"[DEBUG] docx文件格式验证通过: {docx_path}")
            
            # 输出文档基本信息用于调试
            logger.info(f"[DEBUG] 文档包含 {len(doc.sections)} 个section")
            logger.info(f"[DEBUG] 文档包含 {len(doc.paragraphs)} 个段落")
            logger.info(f"[DEBUG] 文档包含 {len(doc.tables)} 个表格")
            
            # 检查页面设置
            if len(doc.sections) > 0:
                section = doc.sections[0]
                logger.info(f"[DEBUG] 页面设置 - 宽度: {section.page_width}, 高度: {section.page_height}")
                logger.info(f"[DEBUG] 页面边距 - 上: {section.top_margin}, 下: {section.bottom_margin}")
                logger.info(f"[DEBUG] 页面边距 - 左: {section.left_margin}, 右: {section.right_margin}")
            
            return True
            
        except Exception as e:
            logger.error(f"[DEBUG] docx文件格式验证失败: {e}", exc_info=True)
            raise RuntimeError(f"Word文档格式错误，无法保证PDF转换质量: {e}")




class WeChatTemplateService:
    """微信模板生成服务 - 移植自GUI项目的modify_wechat_template_info.py"""

    def __init__(self):
        self.template_formats = {
            'standard': self._format_standard_template,
            'detailed': self._format_detailed_template,
            'simple': self._format_simple_template
        }

    def generate_wechat_template(self, report_data: Dict[str, Any], template_type: str = 'standard') -> str:
        """
        生成微信模板内容

        Args:
            report_data: 报告数据
            template_type: 模板类型 ('standard', 'detailed', 'simple')

        Returns:
            str: 格式化的微信模板内容
        """
        try:
            formatter = self.template_formats.get(template_type, self._format_standard_template)
            template_content = formatter(report_data)

            return template_content

        except Exception as e:
            logger.error(f"微信模板生成失败: {e}")
            return self._generate_error_template(str(e))

    def _format_standard_template(self, report_data: Dict[str, Any]) -> str:
        """标准微信模板格式"""
        contact_person = report_data.get('contact_person', '')
        project_address = report_data.get('project_address', '')
        phone = report_data.get('phone', '')
        check_type_display = report_data.get('check_type_display', '初检')
        points_data = report_data.get('points_data', [])

        # 构建点位结果摘要
        results_summary = self._build_results_summary(points_data)

        template = f"""🏠 室内空气检测报告

👤 客户信息：
联系人：{contact_person}
地址：{project_address}
电话：{phone}

🔬 检测信息：
检测类型：{check_type_display}
检测项目：甲醛
检测标准：GB/T 18883-2002

📊 检测结果：
{results_summary}

📋 报告说明：
• 检测依据：GB/T 18883-2002《室内空气质量标准》
• 甲醛标准限值：≤0.08 mg/m³
• 报告已生成，请查收

如有疑问，请随时联系我们！"""

        return template.strip()

    def _format_detailed_template(self, report_data: Dict[str, Any]) -> str:
        """详细微信模板格式"""
        contact_person = report_data.get('contact_person', '')
        project_address = report_data.get('project_address', '')
        phone = report_data.get('phone', '')
        sampling_date = report_data.get('sampling_date', '')
        temperature = report_data.get('temperature', '')
        humidity = report_data.get('humidity', '')
        check_type_display = report_data.get('check_type_display', '初检')
        points_data = report_data.get('points_data', [])

        # 构建详细点位信息
        detailed_results = self._build_detailed_results(points_data)

        template = f"""🏠 室内空气质量检测报告（详细版）

👤 客户信息：
• 联系人：{contact_person}
• 项目地址：{project_address}
• 联系电话：{phone}

🔬 检测详情：
• 采样日期：{sampling_date}
• 现场温度：{temperature}℃
• 现场湿度：{humidity}%
• 检测类型：{check_type_display}
• 检测项目：甲醛

📊 详细检测结果：
{detailed_results}

📋 技术说明：
• 检测依据：GB/T 18883-2002《室内空气质量标准》
• 甲醛标准限值：≤0.08 mg/m³
• 检测方法：分光光度法
• 采样时间：45分钟

✅ 报告状态：已完成
如有任何疑问，请随时联系我们！"""

        return template.strip()

    def _format_simple_template(self, report_data: Dict[str, Any]) -> str:
        """简单微信模板格式"""
        contact_person = report_data.get('contact_person', '')
        project_address = report_data.get('project_address', '')
        check_type_display = report_data.get('check_type_display', '初检')
        points_data = report_data.get('points_data', [])

        # 获取检测结果概况
        total_points = len(points_data)
        qualified_count = self._count_qualified_points(points_data)

        template = f"""🏠 检测报告通知

{contact_person}，您好！
{project_address}的{check_type_display}报告已完成。

检测概况：
• 检测点位：{total_points}个
• 合格点位：{qualified_count}个
• 检测项目：甲醛

报告已生成，请查收！
如有疑问请联系我们。"""

        return template.strip()

    def _build_results_summary(self, points_data: List[Tuple[str, str]]) -> str:
        """构建检测结果摘要"""
        if not points_data:
            return "暂无检测数据"

        summary_lines = []
        for i, (point, value) in enumerate(points_data[:5]):  # 最多显示5个点位
            try:
                value_float = float(value)
                status = "✅ 合格" if value_float <= 0.08 else "❌ 超标"
                summary_lines.append(f"• {point}：{value} mg/m³ {status}")
            except (ValueError, TypeError):
                summary_lines.append(f"• {point}：{value}")

        if len(points_data) > 5:
            summary_lines.append(f"... 等共{len(points_data)}个检测点位")

        return "\n".join(summary_lines)

    def _build_detailed_results(self, points_data: List[Tuple[str, str]]) -> str:
        """构建详细检测结果"""
        if not points_data:
            return "暂无检测数据"

        result_lines = []
        for i, (point, value) in enumerate(points_data):
            try:
                value_float = float(value)
                status = "合格" if value_float <= 0.08 else "超标"
                result_lines.append(f"{i+1:2d}. {point:<8} {value:>6} mg/m³ ({status})")
            except (ValueError, TypeError):
                result_lines.append(f"{i+1:2d}. {point:<8} {value}")

        return "\n".join(result_lines)

    def _count_qualified_points(self, points_data: List[Tuple[str, str]]) -> int:
        """统计合格点位数量"""
        qualified_count = 0
        for _, value in points_data:
            try:
                value_float = float(value)
                if value_float <= 0.08:
                    qualified_count += 1
            except (ValueError, TypeError):
                continue
        return qualified_count

    def _generate_error_template(self, error_msg: str) -> str:
        """生成错误模板"""
        return f"""❌ 微信模板生成失败

错误信息：{error_msg}

请联系技术支持解决此问题。
生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"""
