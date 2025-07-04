"""
报告生成异步任务
"""
import os
import time
import logging
from datetime import datetime
from celery import shared_task
from django.utils import timezone
from django.conf import settings
from django.core.files.base import ContentFile
from django.db import transaction
from .models import Report, ReportTemplate

logger = logging.getLogger(__name__)


@shared_task(bind=True, max_retries=3)
@transaction.atomic
def generate_report(self, report_id, user_id, template_id=None):
    """
    生成报告任务

    Args:
        report_id: 报告ID
        user_id: 用户ID
        template_id: 模板ID（可选）
    """
    report = None
    try:
        # 获取报告对象
        report = Report.objects.get(id=report_id)
        logger.info(f"开始生成报告: report_{report.id}")
        
        # 更新生成状态
        report.generation_started_at = timezone.now()
        report.is_generated = False
        report.error_message = ''
        report.save()
        logger.info(f"报告状态已重置: report_{report.id}")
        
        # 获取模板
        template = None
        if template_id:
            try:
                template = ReportTemplate.objects.get(id=template_id, is_active=True)
                logger.info(f"使用模板: {template.name}")
            except ReportTemplate.DoesNotExist:
                logger.warning(f"模板不存在: {template_id}")
                pass
        
        # 生成Word文档
        logger.info("开始生成Word文档...")
        docx_content = generate_word_document(report, template)
        logger.info(f"Word文档生成成功，大小: {len(docx_content)} bytes")
        
        # 保存Word文档
        docx_filename = f"report_{report.id}_{int(time.time())}.docx"
        report.docx_file.save(
            docx_filename,
            ContentFile(docx_content),
            save=False
        )
        logger.info(f"Word文档已保存: {docx_filename}")
        
        # 生成PDF文档
        logger.info("开始生成PDF文档...")
        pdf_content = convert_to_pdf_sync(docx_content)
        logger.info(f"PDF文档生成成功，大小: {len(pdf_content)} bytes")
        
        # 保存PDF文档
        pdf_filename = f"report_{report.id}_{int(time.time())}.pdf"
        report.pdf_file.save(
            pdf_filename,
            ContentFile(pdf_content),
            save=False
        )
        logger.info(f"PDF文档已保存: {pdf_filename}")

        # 更新生成状态
        report.is_generated = True
        report.generation_completed_at = timezone.now()

        # 先保存PDF文件和状态
        report.save()
        logger.info(f"报告状态已更新: report_{report.id}, is_generated={report.is_generated}")

        # 删除原始Word文件（如果设置了删除选项）
        if report.delete_original_docx and report.docx_file:
            try:
                # 删除物理文件
                if os.path.exists(report.docx_file.path):
                    os.remove(report.docx_file.path)
                    logger.info(f"已删除物理Word文件: {report.docx_file.path}")
                # 清空文件字段
                report.docx_file = None
                # 再次保存以更新文件字段
                report.save()
                logger.info(f"已清空Word文件字段: report_{report.id}")
            except Exception as e:
                logger.warning(f"删除原始Word文件失败: {e}")

        # 确保最终状态正确
        report.refresh_from_db()
        logger.info(f"最终报告状态: report_{report.id}, is_generated={report.is_generated}, pdf_file={bool(report.pdf_file)}")
        
        return {
            'status': 'success',
            'report_id': report.id,
            'docx_file': report.docx_file.name if report.docx_file else None,
            'pdf_file': report.pdf_file.name if report.pdf_file else None,
            'generation_time': (report.generation_completed_at - report.generation_started_at).total_seconds()
        }
        
    except Exception as e:
        error_message = f"报告生成失败: {str(e)}"
        logger.error(f"报告生成异常: {error_message}", exc_info=True)
        
        # 更新错误状态
        if report is not None:
            try:
                report.error_message = error_message
                report.generation_completed_at = timezone.now()
                report.is_generated = False  # 确保失败时状态为False
                report.save()
                logger.info(f"错误状态已更新: report_{report.id}")
            except Exception as save_error:
                logger.error(f"保存错误状态失败: {save_error}")
        
        # 重试机制
        if self.request.retries < self.max_retries:
            logger.info(f"准备重试，当前重试次数: {self.request.retries}")
            raise self.retry(countdown=60 * (self.request.retries + 1))
        
        return {
            'status': 'error',
            'error': error_message,
            'report_id': report_id if report else None
        }


def generate_word_document(report, template=None):
    """
    生成Word文档 - 使用真实的报告生成服务

    Args:
        report: 报告对象
        template: 模板对象（可选）

    Returns:
        bytes: Word文档内容
    """
    try:
        from .services import ReportGenerationService

        # 创建报告生成服务
        report_service = ReportGenerationService()

        # 准备OCR结果数据，处理空值情况
        ocr_result_data = {}
        if report.ocr_result:
            ocr_result_data = {
                'phone': getattr(report.ocr_result, 'phone', '') or '',
                'date': getattr(report.ocr_result, 'date', '') or '',
                'temperature': getattr(report.ocr_result, 'temperature', '') or '',
                'humidity': getattr(report.ocr_result, 'humidity', '') or '',
                'check_type': getattr(report.ocr_result, 'check_type', 'initial') or 'initial',
                'points_data': getattr(report.ocr_result, 'points_data', {}) or {}
            }
        else:
            # 如果没有OCR结果，使用默认值
            logger.warning(f"报告 {report.id} 没有关联的OCR结果，使用默认值")
            ocr_result_data = {
                'phone': '',
                'date': '',
                'temperature': '',
                'humidity': '',
                'check_type': 'initial',
                'points_data': {}
            }

        # 确保form_data不为空
        form_data = report.form_data or {}

        logger.info(f"准备生成报告，OCR数据: {ocr_result_data}, 表单数据: {form_data}")

        # 生成报告
        docx_content, pdf_content = report_service.generate_report(
            ocr_result_data,
            form_data
        )

        return docx_content

    except Exception as e:
        logger.error(f"Word文档生成失败: {e}", exc_info=True)
        
        # 如果生成失败，创建一个简单的降级文档
        try:
            from docx import Document
            from io import BytesIO
            
            # 创建简单的报告文档
            doc = Document()
            doc.add_heading('室内空气质量检测报告', 0)
            
            # 添加基本信息
            doc.add_paragraph(f'报告标题: {report.title}')
            doc.add_paragraph(f'报告类型: {report.get_report_type_display()}')
            doc.add_paragraph(f'生成时间: {timezone.now().strftime("%Y-%m-%d %H:%M:%S")}')
            
            # 添加表单数据
            form_data = report.form_data or {}
            if form_data:
                doc.add_heading('报告信息', level=1)
                for key, value in form_data.items():
                    if value:
                        doc.add_paragraph(f'{key}: {value}')
            
            # 添加错误信息
            doc.add_heading('生成说明', level=1)
            doc.add_paragraph(f'注意：由于生成过程中遇到问题，此为降级版本报告。')
            doc.add_paragraph(f'错误信息：{str(e)}')
            doc.add_paragraph(f'如需完整报告，请联系系统管理员。')
            
            # 保存到内存
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            logger.info("已生成降级版本的Word文档")
            return doc_buffer.getvalue()
            
        except Exception as fallback_error:
            logger.error(f"降级文档生成也失败: {fallback_error}")
            # 最后的降级方案：纯文本
            content = f"""
报告标题: {report.title}
报告类型: {report.get_report_type_display()}
生成时间: {timezone.now().strftime('%Y-%m-%d %H:%M:%S')}

生成错误: {str(e)}
降级错误: {str(fallback_error)}

请联系系统管理员解决此问题。
            """.strip()
            return content.encode('utf-8')


def convert_to_pdf_sync(docx_content):
    """
    同步转换Word文档为PDF，带超时机制
    
    只使用能保持完整格式的转换方法，绝不提供降级方案。
    格式完整性是报告质量的基本要求。

    Args:
        docx_content: Word文档内容

    Returns:
        bytes: PDF文档内容
        
    Raises:
        RuntimeError: 当无法保证格式完整性时
    """
    import signal
    import threading
    from concurrent.futures import ThreadPoolExecutor, TimeoutError as FutureTimeoutError

    def _convert_with_timeout():
        """在线程中执行PDF转换"""
        try:
            from .services import ReportGenerationService

            # 创建报告生成服务
            report_service = ReportGenerationService()

            # 调用PDF转换方法
            pdf_content = report_service._convert_to_pdf(docx_content)

            return pdf_content

        except Exception as e:
            logger.error(f"PDF转换失败: {e}")
            raise e

    try:
        # 使用线程池执行转换，设置60秒超时
        with ThreadPoolExecutor(max_workers=1) as executor:
            future = executor.submit(_convert_with_timeout)
            try:
                pdf_content = future.result(timeout=60)  # 60秒超时
                return pdf_content
            except FutureTimeoutError:
                logger.warning("PDF转换超时，使用降级方案")
                raise TimeoutError("PDF转换超时")

    except Exception as e:
        logger.error(f"PDF转换失败: {e}")
        # 绝对不提供降级方案，直接抛出异常
        # 格式完整性是报告质量的基本要求
        raise RuntimeError(f"PDF转换失败，无法保证格式完整性: {e}")


@shared_task(bind=True, max_retries=3)
def convert_to_pdf(self, docx_content):
    """
    将Word文档转换为PDF - 使用真实的PDF转换服务

    Args:
        docx_content: Word文档内容

    Returns:
        bytes: PDF文档内容
    """
    try:
        from .services import ReportGenerationService

        # 创建报告生成服务
        report_service = ReportGenerationService()

        # 调用PDF转换方法
        pdf_content = report_service._convert_to_pdf(docx_content)

        return pdf_content

    except Exception as e:
        # 重试机制
        if self.request.retries < self.max_retries:
            raise self.retry(countdown=30 * (self.request.retries + 1))

        raise e


@shared_task
def generate_wechat_template(report_id, template_type='standard'):
    """
    生成微信模板任务

    Args:
        report_id: 报告ID
        template_type: 模板类型

    Returns:
        dict: 微信模板内容
    """
    try:
        from .services import WeChatTemplateService

        # 获取报告对象
        report = Report.objects.get(id=report_id)

        # 准备报告数据
        report_data = {
            'contact_person': report.form_data.get('contact_person', ''),
            'project_address': report.form_data.get('project_address', ''),
            'phone': report.ocr_result.phone if report.ocr_result else '',
            'sampling_date': report.form_data.get('sampling_date', ''),
            'temperature': report.form_data.get('temperature', ''),
            'humidity': report.form_data.get('humidity', ''),
            'check_type_display': '初检' if report.form_data.get('check_type') == 'initial' else '复检',
            'points_data': []
        }

        # 转换点位数据格式
        if report.ocr_result and report.ocr_result.points_data:
            points_data = []
            for point_name, point_value in report.ocr_result.points_data.items():
                points_data.append((point_name, str(point_value)))
            report_data['points_data'] = points_data

        # 创建微信模板服务
        wechat_service = WeChatTemplateService()

        # 生成模板内容
        template_content = wechat_service.generate_wechat_template(report_data, template_type)

        return {
            'status': 'success',
            'report_id': report_id,
            'template_type': template_type,
            'template_content': template_content,
            'generated_at': timezone.now().isoformat()
        }

    except Report.DoesNotExist:
        return {
            'status': 'error',
            'error': '报告不存在',
            'report_id': report_id
        }
    except Exception as e:
        return {
            'status': 'error',
            'error': str(e),
            'report_id': report_id
        }


@shared_task
def batch_generate_reports(report_ids, user_id, template_id=None):
    """
    批量生成报告
    
    Args:
        report_ids: 报告ID列表
        user_id: 用户ID
        template_id: 模板ID（可选）
    """
    results = []
    
    for report_id in report_ids:
        try:
            result = generate_report.delay(report_id, user_id, template_id)
            results.append({
                'report_id': report_id,
                'task_id': result.id,
                'status': 'started'
            })
        except Exception as e:
            results.append({
                'report_id': report_id,
                'status': 'error',
                'error': str(e)
            })
    
    return {
        'total_reports': len(report_ids),
        'started_tasks': len([r for r in results if r['status'] == 'started']),
        'failed_tasks': len([r for r in results if r['status'] == 'error']),
        'results': results
    }


@shared_task
def cleanup_old_reports():
    """清理旧的报告文件"""
    from datetime import timedelta
    
    # 删除30天前的报告文件
    cutoff_date = timezone.now() - timedelta(days=30)
    old_reports = Report.objects.filter(
        created_at__lt=cutoff_date,
        is_generated=True
    )
    
    cleaned_count = 0
    for report in old_reports:
        try:
            # 删除文件
            if report.docx_file and os.path.exists(report.docx_file.path):
                os.remove(report.docx_file.path)
            if report.pdf_file and os.path.exists(report.pdf_file.path):
                os.remove(report.pdf_file.path)
            
            # 重置文件字段
            report.docx_file = None
            report.pdf_file = None
            report.is_generated = False
            report.save()
            
            cleaned_count += 1
            
        except Exception as e:
            print(f"清理报告 {report.id} 失败: {e}")
    
    return f"清理了 {cleaned_count} 个旧报告文件"


@shared_task
def generate_template_preview(template_id, sample_data):
    """
    生成模板预览
    
    Args:
        template_id: 模板ID
        sample_data: 示例数据
        
    Returns:
        dict: 预览结果
    """
    try:
        template = ReportTemplate.objects.get(id=template_id, is_active=True)
        
        # TODO: 实现模板预览生成逻辑
        time.sleep(1)  # 模拟生成时间
        
        return {
            'status': 'success',
            'template_id': template_id,
            'preview_url': f'/media/previews/template_{template_id}_preview.pdf',
            'generated_at': timezone.now().isoformat()
        }
        
    except ReportTemplate.DoesNotExist:
        return {
            'status': 'error',
            'error': '模板不存在'
        }
    except Exception as e:
        return {
            'status': 'error',
            'error': str(e)
        }
